import fitz  # PyMuPDF
import docx
from pathlib import Path
from typing import Dict, List, Optional
import logging
import re
from datetime import datetime

class DocumentProcessor:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.date_patterns = [
            r'\d{1,2}/\d{1,2}/\d{4}',  # DD/MM/YYYY
            r'\d{4}-\d{2}-\d{2}',      # YYYY-MM-DD
            r'\d{1,2}\.\d{1,2}\.\d{4}', # DD.MM.YYYY
            r'\d{4}/\d{2}/\d{2}'       # YYYY/MM/DD
        ]
    
    def extract_from_pdf(self, pdf_path: str) -> Dict:
        """Estrae testo e tabelle da un documento PDF."""
        try:
            doc = fitz.open(pdf_path)
            text = ""
            tables = []
            
            for page in doc:
                text += page.get_text()
                # Estrai tabelle (implementazione base)
                tables.extend(self._extract_tables_from_page(page))
            
            return {
                'text': text,
                'tables': tables
            }
        except Exception as e:
            self.logger.error(f"Errore nell'estrazione da PDF {pdf_path}: {str(e)}")
            return {'text': '', 'tables': []}
    
    def extract_from_word(self, docx_path: str) -> Dict:
        """Estrae testo e tabelle da un documento Word."""
        try:
            doc = docx.Document(docx_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            tables = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)
            
            return {
                'text': text,
                'tables': tables
            }
        except Exception as e:
            self.logger.error(f"Errore nell'estrazione da Word {docx_path}: {str(e)}")
            return {'text': '', 'tables': []}
    
    def _extract_tables_from_page(self, page) -> List:
        """Estrae tabelle da una pagina PDF."""
        tables = []
        try:
            # Estrai il testo della pagina
            text = page.get_text()
            lines = text.split('\n')
            
            # Cerca pattern di tabelle (righe con più di 2 colonne)
            table_lines = []
            current_table = []
            
            for line in lines:
                # Rimuovi spazi extra e caratteri speciali
                line = re.sub(r'\s+', ' ', line).strip()
                
                # Se la riga ha più di 2 colonne, potrebbe essere parte di una tabella
                if len(line.split()) > 2:
                    current_table.append(line.split())
                elif current_table:
                    # Se troviamo una riga vuota o con poche colonne dopo una tabella,
                    # salviamo la tabella corrente
                    tables.append(current_table)
                    current_table = []
            
            # Aggiungi l'ultima tabella se presente
            if current_table:
                tables.append(current_table)
            
        except Exception as e:
            self.logger.error(f"Errore nell'estrazione delle tabelle: {str(e)}")
        
        return tables
    
    def _parse_date(self, date_str: str) -> Optional[datetime]:
        """Converte una stringa di data in un oggetto datetime."""
        if not date_str:
            return None
            
        for pattern in self.date_patterns:
            match = re.search(pattern, date_str)
            if match:
                date_str = match.group(0)
                try:
                    # Prova diversi formati
                    for fmt in ['%d/%m/%Y', '%Y-%m-%d', '%d.%m.%Y', '%Y/%m/%d']:
                        try:
                            return datetime.strptime(date_str, fmt)
                        except ValueError:
                            continue
                except Exception:
                    continue
        return None
    
    def extract_mission_data(self, text: str) -> Dict:
        """Estrae dati specifici delle missioni dal testo."""
        data = {}
        
        # Pattern per estrarre informazioni
        patterns = {
            'budget': r'budget[:\s]+€?\s*([\d,.]+)',
            'personnel': r'personnel[:\s]+(\d+)',
            'start_date': r'start(?:ing)?\s*date[:\s]+([^\n]+)',
            'end_date': r'end(?:ing)?\s*date[:\s]+([^\n]+)',
            'location': r'location[:\s]+([^\n]+)',
            'mandate': r'mandate[:\s]+([^\n]+)',
            'mission_name': r'mission\s*(?:name)?[:\s]+([^\n]+)',
            'country': r'country[:\s]+([^\n]+)',
            'type': r'type[:\s]+([^\n]+)'
        }
        
        # Cerca i pattern nel testo
        for key, pattern in patterns.items():
            matches = re.finditer(pattern, text.lower())
            for match in matches:
                value = match.group(1).strip()
                
                # Converti le date
                if 'date' in key:
                    parsed_date = self._parse_date(value)
                    if parsed_date:
                        data[key] = parsed_date.strftime('%Y-%m-%d')
                # Converti i numeri
                elif key in ['budget', 'personnel']:
                    try:
                        if key == 'budget':
                            # Rimuovi eventuali simboli di valuta e spazi
                            value = re.sub(r'[^\d.,]', '', value)
                            # Sostituisci la virgola con il punto per i decimali
                            value = value.replace(',', '.')
                        data[key] = float(value)
                    except ValueError:
                        self.logger.warning(f"Impossibile convertire {value} in numero per {key}")
                else:
                    data[key] = value
        
        return data
    
    def process_document(self, file_path: str) -> Dict:
        """Processa un documento (PDF o Word) ed estrae i dati rilevanti."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            self.logger.error(f"File non trovato: {file_path}")
            return {}
        
        try:
            if file_path.suffix.lower() == '.pdf':
                content = self.extract_from_pdf(str(file_path))
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                content = self.extract_from_word(str(file_path))
            else:
                self.logger.warning(f"Formato file non supportato: {file_path}")
                return {}
            
            # Estrai dati specifici delle missioni
            mission_data = self.extract_mission_data(content['text'])
            
            # Estrai dati dalle tabelle
            table_data = self._extract_data_from_tables(content['tables'])
            mission_data.update(table_data)
            
            return {
                'mission_data': mission_data,
                'tables': content['tables']
            }
            
        except Exception as e:
            self.logger.error(f"Errore nel processare il documento {file_path}: {str(e)}")
            return {}
    
    def _extract_data_from_tables(self, tables: List) -> Dict:
        """Estrae dati rilevanti dalle tabelle."""
        data = {}
        
        for table in tables:
            # Cerca intestazioni rilevanti
            headers = table[0] if table else []
            for i, header in enumerate(headers):
                header = header.lower()
                if 'budget' in header or 'cost' in header:
                    # Cerca valori numerici nella colonna
                    for row in table[1:]:
                        if i < len(row):
                            value = row[i]
                            try:
                                value = float(re.sub(r'[^\d.,]', '', value).replace(',', '.'))
                                data['budget'] = value
                                break
                            except ValueError:
                                continue
                elif 'personnel' in header or 'staff' in header:
                    # Cerca valori numerici nella colonna
                    for row in table[1:]:
                        if i < len(row):
                            value = row[i]
                            try:
                                value = int(re.sub(r'[^\d]', '', value))
                                data['personnel'] = value
                                break
                            except ValueError:
                                continue
        
        return data 