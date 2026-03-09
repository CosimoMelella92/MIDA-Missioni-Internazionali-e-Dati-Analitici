"""
DOCX Parser Module
Base functionality for Word (.docx) text and table extraction
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional

import docx


class DOCXParser:
    """Base DOCX parser for text and table extraction"""
    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def extract_text_from_docx(self, docx_path: str) -> Dict[str, any]:
        """
        Extract text and tables from a DOCX file
        Args:
            docx_path: Path to DOCX file
        Returns:
            Dictionary with extracted text and tables
        """
        docx_path = Path(docx_path)
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        self.logger.info(f"Processing DOCX: {docx_path.name}")
        try:
            doc = docx.Document(str(docx_path))
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = '\n'.join(paragraphs)
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append({
                        'data': table_data,
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data else 0
                    })
            return {
                'filename': docx_path.name,
                'full_text': full_text,
                'tables': tables,
                'paragraphs': paragraphs
            }
        except Exception as e:
            self.logger.error(f"Error processing DOCX {docx_path}: {str(e)}")
            raise

    def process_docx_directory(self, directory_path: str) -> List[Dict]:
        """
        Process all DOCX files in a directory
        Args:
            directory_path: Path to directory containing DOCX files
        Returns:
            List of extracted data from all DOCX files
        """
        directory = Path(directory_path)
        results = []
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        docx_files = list(directory.glob("*.docx"))
        self.logger.info(f"Found {len(docx_files)} DOCX files in {directory_path}")
        for docx_file in docx_files:
            try:
                extracted_data = self.extract_text_from_docx(str(docx_file))
                results.append({
                    'file': docx_file.name,
                    'extracted_data': extracted_data
                })
            except Exception as e:
                self.logger.error(f"Error processing {docx_file.name}: {str(e)}")
                results.append({
                    'file': docx_file.name,
                    'error': str(e)
                })
        return results
